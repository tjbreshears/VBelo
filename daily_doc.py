from docx import Document
from datetime import datetime, timedelta
import time
import csv

import config


games = []
with open("outputs/games_output.csv", 'r') as data:
    for line in csv.DictReader(data):
        games.append(line)

teams = []
with open("inputs/VBelo - teams.csv", 'r') as data:
    for line in csv.DictReader(data):
        teams.append(line)
for i in range(len(teams)):
    teams[i]['elo'] = int(teams[i]['elo'])

rankings = []
with open(f"inputs/polls/avca_{config.current_year}.csv", 'r') as data:
    for line in csv.DictReader(data):
        rankings.append(line)


# day is the day sent (not created)
def create_doc(day):
    bwc = []
    cc = []
    ecc = []
    eiva = []
    iva = []
    miva = []
    mpsf = []
    nec = []
    siac = []
    nc = []
    away = ''
    home = ''
    away_c = ''
    home_c = ''
    away_p = ''
    home_p = ''
    away_r = ''
    home_r = ''
    away_s = ''
    home_s = ''
    away_p_1 = ''
    away_p_2 = ''
    away_p_3 = ''
    away_p_4 = ''
    away_p_5 = ''
    home_p_1 = ''
    home_p_2 = ''
    home_p_3 = ''
    home_p_4 = ''
    home_p_5 = ''

    document = Document()
    long_date = datetime.strptime(day, '%m/%d/%Y').strftime('%B %d, %Y')
    yesterday = (datetime.strptime(day, '%m/%d/%Y') + timedelta(days=-1)).strftime('%#m/%#d/%Y')

    document.add_heading(f'VBelo Daily: {long_date}', level=1)
    document.add_paragraph('Subscribe to the VBelo Report to get fresh volleyball content float served '
                           'straight into your inbox.')

# section needs to be added
    document.add_heading("Yesterday's Scores", level=2)
    for game in games:
        if yesterday in game['date']:
            for team in teams:
                if team['short_name'] == game['t1']:
                    away = team['full_name']
                    for rank in rankings:
                        if rank["latest"] == team['short_name']:
                            away = "#" + str(rank["rank"]) + " " + away
                    away_c = team['conference']
                    away_p = round(float(game['probability_team1'])*100)
                    away_r = game['r_t1']
                    away_s = game['s_t1']
                    away_p_1 = game['s1_t1']
                    away_p_2 = game['s2_t1']
                    away_p_3 = game['s3_t1']
                    away_p_4 = game['s4_t1']
                    away_p_5 = game['s5_t1']
                if team['short_name'] == game['t2']:
                    home = team['full_name']
                    for rank in rankings:
                        if rank["latest"] == team['short_name']:
                            home = "#" + str(rank["rank"]) + " " + home
                    home_c = team['conference']
                    home_p = round(float(game['probability_team2'])*100)
                    home_r = game['r_t2']
                    home_s = game['s_t2']
                    home_p_1 = game['s1_t2']
                    home_p_2 = game['s2_t2']
                    home_p_3 = game['s3_t2']
                    home_p_4 = game['s4_t2']
                    home_p_5 = game['s5_t2']
            match_data = {'away': away, 'away_p': away_p, 'away_s': away_s, 'home': home, 'home_p': home_p,
                          'away_r': away_r, 'home_s': home_s, 'home_r': home_r, 'away_p_1': away_p_1,
                          'away_p_2': away_p_2, 'away_p_3': away_p_3, 'away_p_4': away_p_4, 'away_p_5': away_p_5,
                          'home_p_1': home_p_1, 'home_p_2': home_p_2, 'home_p_3': home_p_3, 'home_p_4': home_p_4,
                          'home_p_5': home_p_5}

            if away_c == 'BWC' and home_c == 'BWC':
                bwc.append(match_data)
            elif away_c == 'CC' and home_c == 'CC':
                cc.append(match_data)
            elif away_c == 'ECC' and home_c == 'ECC':
                ecc.append(match_data)
            elif away_c == 'EIVA' and home_c == 'EIVA':
                eiva.append(match_data)
            elif away_c == 'IVA' and home_c == 'IVA':
                iva.append(match_data)
            elif away_c == 'MIVA' and home_c == 'MIVA':
                miva.append(match_data)
            elif away_c == 'MPSF' and home_c == 'MPSF':
                mpsf.append(match_data)
            elif away_c == 'NEC' and home_c == 'NEC':
                nec.append(match_data)
            elif away_c == 'SIAC' and home_c == 'SIAC':
                siac.append(match_data)
            elif away_c != home_c or (away_c == 'IND' and home_c == 'IND'):
                nc.append(match_data)

    if len(bwc) != 0:
        document.add_heading('Big West', level=4)
        for match in bwc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        bwc = []
    if len(cc) != 0:
        document.add_heading('ConfCarolinas', level=4)
        for match in cc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        cc = []
    if len(ecc) != 0:
        document.add_heading('ECC', level=4)
        for match in ecc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        ecc = []
    if len(eiva) != 0:
        document.add_heading('EIVA', level=4)
        for match in eiva:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        eiva = []
    if len(iva) != 0:
        document.add_heading('IVA', level=4)
        for match in iva:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        iva = []
    if len(miva) != 0:
        document.add_heading('MIVA', level=4)
        for match in miva:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        miva = []
    if len(mpsf) != 0:
        document.add_heading('MPSF', level=4)
        for match in mpsf:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        mpsf = []
    if len(nec) != 0:
        document.add_heading('NEC', level=4)
        for match in nec:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        nec = []
    if len(siac) != 0:
        document.add_heading('SIAC', level=4)
        for match in siac:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        siac = []
    if len(nc) != 0:
        document.add_heading('Non-conference', level=4)
        for match in nc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            team_1_s = match['away_s']
            team_2_s = match['home_s']
            team_1_p_1 = match['away_p_1']
            team_1_p_2 = match['away_p_2']
            team_1_p_3 = match['away_p_3']
            team_1_p_4 = match['away_p_4']
            team_1_p_5 = match['away_p_5']
            team_2_p_1 = match['home_p_1']
            team_2_p_2 = match['home_p_2']
            team_2_p_3 = match['home_p_3']
            team_2_p_4 = match['home_p_4']
            team_2_p_5 = match['home_p_5']
            if match['away_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_1} ({team_1_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_2} ({team_2_p}%)')
                run.add_break()
                p.add_run(f'{team_1_s}-{team_2_s}').bold = True
                p.add_run(f' ({team_1_p_1}-{team_2_p_1}, {team_1_p_2}-{team_2_p_2}, {team_1_p_3}-{team_2_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_1_p_4}-{team_2_p_4}, {team_1_p_5}-{team_2_p_5})')
            if match['home_r'] == '1':
                p = document.add_paragraph('')
                p.add_run(f'{team_2} ({team_2_p}%) ').bold = True
                p.add_run('def.').italic = True
                run = p.add_run(f' {team_1} ({team_1_p}%)')
                run.add_break()
                p.add_run(f'{team_2_s}-{team_1_s}').bold = True
                p.add_run(f' ({team_2_p_1}-{team_1_p_1}, {team_2_p_2}-{team_1_p_2}, {team_2_p_3}-{team_1_p_3}')
                if int(team_1_s) + int(team_2_s) == 3:
                    p.add_run(')')
                elif int(team_1_s) + int(team_2_s) == 4:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4})')
                elif int(team_1_s) + int(team_2_s) == 5:
                    p.add_run(f', {team_2_p_4}-{team_1_p_4}, {team_2_p_5}-{team_1_p_5})')
        nc = []

    document.add_heading("Today's Match Projections", level=2)
    for game in games:
        if day in game['date']:
            for team in teams:
                if team['short_name'] == game['t1']:
                    away = team['full_name']
                    for rank in rankings:
                        if rank["latest"] == team['short_name']:
                            away = "#" + str(rank["rank"]) + " " + away
                    away_c = team['conference']
                    away_p = round(float(game['probability_team1'])*100)
                if team['short_name'] == game['t2']:
                    home = team['full_name']
                    for rank in rankings:
                        if rank["latest"] == team['short_name']:
                            home = "#" + str(rank["rank"]) + " " + home
                    home_c = team['conference']
                    home_p = round(float(game['probability_team2'])*100)
            match_data = {'away': away, 'away_p': away_p, 'home': home, 'home_p': home_p}
            if away_c != home_c or (away_c == 'IND' and home_c == 'IND'):
                nc.append(match_data)
            elif away_c == 'BWC' and home_c == 'BWC':
                bwc.append(match_data)
            elif away_c == 'CC' and home_c == 'CC':
                cc.append(match_data)
            elif away_c == 'ECC' and home_c == 'ECC':
                ecc.append(match_data)
            elif away_c == 'EIVA' and home_c == 'EIVA':
                eiva.append(match_data)
            elif away_c == 'IVA' and home_c == 'IVA':
                iva.append(match_data)
            elif away_c == 'MIVA' and home_c == 'MIVA':
                miva.append(match_data)
            elif away_c == 'MPSF' and home_c == 'MPSF':
                mpsf.append(match_data)
            elif away_c == 'NEC' and home_c == 'NEC':
                nec.append(match_data)
            elif away_c == 'SIAC' and home_c == 'SIAC':
                siac.append(match_data)

    if len(bwc) != 0:
        document.add_heading('Big West', level=4)
        for match in bwc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(cc) != 0:
        document.add_heading('ConfCarolinas', level=4)
        for match in cc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(ecc) != 0:
        document.add_heading('ECC', level=4)
        for match in ecc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(eiva) != 0:
        document.add_heading('EIVA', level=4)
        for match in eiva:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(iva) != 0:
        document.add_heading('IVA', level=4)
        for match in iva:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(miva) != 0:
        document.add_heading('MIVA', level=4)
        for match in miva:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(mpsf) != 0:
        document.add_heading('MPSF', level=4)
        for match in mpsf:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(nec) != 0:
        document.add_heading('NEC', level=4)
        for match in nec:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(siac) != 0:
        document.add_heading('SIAC', level=4)
        for match in siac:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    if len(nc) != 0:
        document.add_heading('Non-conference', level=4)
        for match in nc:
            team_1 = match['away']
            team_2 = match['home']
            team_1_p = match['away_p']
            team_2_p = match['home_p']
            p = document.add_paragraph(f'{team_1} ({team_1_p}%) ')
            p.add_run('vs').italic = True
            p.add_run(f' {team_2} ({team_2_p}%)')
    document.save('outputs/SR_DATE.docx')


# test case
if __name__ == '__main__':
    start_time = time.time()
    t_now = (datetime.today() + timedelta(days=1)).strftime('%#m/%#d/%Y')
# If testing a specific date, comment out row above and uncomment row below
    # t_now = '3/10/2024'
    create_doc(t_now)
    print("--- %s seconds ---" % (time.time() - start_time))
